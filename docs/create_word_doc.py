"""
Script para convertir el informe Markdown a Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

doc.add_heading("Informe de Seguimiento 1 - Análisis de Algoritmos", 0)

doc.add_heading("Universidad del Quindío", 1)
doc.add_paragraph("Programa de Ingeniería de Sistemas y Computación")
doc.add_paragraph("Curso: Análisis de Algoritmos - 2026-1")

doc.add_heading("1. Introducción y Contexto del Proyecto", 1)

p = doc.add_paragraph()
p.add_run(
    "Este seguimiento 1 hace parte del proyecto final de análisis de algoritmos, "
)
p.add_run("enfocados en el primer punto de este mismo. ").bold = True
p.add_run(
    "Este seguimiento está enfocado no solo en entender cómo funcionan algunos algoritmos de ordenamiento, observar el tiempo de ejecución de estos y entender la complejidad de cada uno, sino también en obtener, entender y organizar toda la información que estos algoritmos reciben. "
)
p.add_run(
    "Uno de los puntos importantes a la hora de usar un algoritmo de ordenamiento es obtener la menor complejidad y tiempo de ejecución posible, con esto optimizar tiempo y operaciones, esto debido a la gran cantidad de datos que se manejan hoy en día en las empresas modernas."
).bold = True

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run(
    "El análisis financiero moderno depende en gran medida de la capacidad computacional para procesar grandes volúmenes de datos históricos y detectar patrones relevantes en el comportamiento de los activos financieros. "
)

doc.add_heading("2. Objetivos del Proyecto", 1)
doc.add_paragraph(
    "El propósito de este proyecto es diseñar e implementar una serie de algoritmos que permitan realizar análisis técnico, estadístico y comparativo de activos financieros, haciendo énfasis en la eficiencia computacional, la correcta fundamentación matemática de las métricas utilizadas y el análisis formal de la complejidad algorítmica."
)

doc.add_paragraph()
doc.add_paragraph("Los objetivos específicos incluyen:", style="List Bullet")

doc.add_heading("3. Activos Financieros Utilizados", 1)
doc.add_paragraph(
    "Para este proyecto se seleccionaron 12 activos financieros que incluyen acciones colombianas y ETFs internacionales."
)

doc.add_heading("3.1 Acciones Colombianas", 2)

table1 = doc.add_table(rows=3, cols=3)
table1.style = "Table Grid"
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = "Símbolo"
hdr_cells[1].text = "Nombre Completo"
hdr_cells[2].text = "Sector"

row1 = table1.rows[1].cells
row1[0].text = "ISA"
row1[1].text = "Interconexión Eléctrica S.A."
row1[2].text = "Energía eléctrica - Transmisión"

row2 = table1.rows[2].cells
row2[0].text = "GEB"
row2[1].text = "Grupo Energía de Bogotá"
row2[2].text = "Energía - Distribución"

doc.add_heading("3.2 ETFs Internacionales", 2)

table2 = doc.add_table(rows=11, cols=3)
table2.style = "Table Grid"
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = "Símbolo"
hdr_cells[1].text = "Nombre Completo"
hdr_cells[2].text = "Descripción"

etfs = [
    ("VOO", "Vanguard S&P 500 ETF", "Replica el índice S&P 500"),
    (
        "VTI",
        "Vanguard Total Stock Market ETF",
        "Todas las empresas del mercado estadounidense",
    ),
    (
        "QQQ",
        "Invesco QQQ Trust",
        "Seguí el Nasdaq-100, concentrando en empresas de tecnología",
    ),
    (
        "SPY",
        "SPDR S&P 500 ETF Trust",
        "Uno de los ETFs más antiguos y líquidos del S&P 500",
    ),
    (
        "VEA",
        "Vanguard FTSE Developed Markets",
        "Mercados desarrollados de Europa, Japón y Australia",
    ),
    (
        "VWO",
        "Vanguard FTSE Emerging Markets",
        "Mercados emergentes como China, India y Brasil",
    ),
    (
        "BND",
        "Vanguard Total Bond Market ETF",
        "Fondo de bonos gubernamentales y corporativos de EE.UU.",
    ),
    (
        "EFA",
        "iShares MSCI EAFE ETF",
        "Países desarrollados de Europa, Asia y Australia",
    ),
    ("EEM", "iShares MSCI Emerging Markets ETF", "Mercados emergentes a nivel mundial"),
    (
        "TLT",
        "iShares 20+ Year Treasury Bond",
        "Bonos del tesoro estadounidenses a largo plazo",
    ),
]

for i, (sym, name, desc) in enumerate(etfs):
    row = table2.rows[i + 1].cells
    row[0].text = sym
    row[1].text = name
    row[2].text = desc

doc.add_heading("4. Arquitectura del Sistema", 1)
doc.add_paragraph(
    "El proyecto sigue una arquitectura de microservicios modular implementada en Python."
)

doc.add_heading("4.1 Estructura de Directorios", 2)
doc.add_paragraph("""src/
├── etl/                      
│   ├── fetcher.py            
│   ├── cleaner.py            
│   └── unifier.py            
├── sorting/                  
│   ├── algorithms.py         
│   ├── comparator.py        
│   └── visualizer.py        
├── services/                 
│   ├── volume_analyzer.py    
│   └── main_runner.py        
└── api/                      
    └── gateway.py""")

doc.add_heading("5. Metodología y Implementación", 1)

doc.add_heading("5.1 Proceso ETL", 2)
doc.add_paragraph("El proceso ETL implementado sigue tres etapas claramente definidas:")

doc.add_paragraph(
    "Extracción de Datos: La obtención de datos financieros se realiza mediante peticiones HTTP directas a la API de Yahoo Finance. Se usó esta implementación manual para cumplir con el requisito de utilizar peticiones explícitas y manejo manual de parsing. Complejidad: O(n × d)."
)

doc.add_paragraph("Limpieza de Datos: Se implementaron tres técnicas de limpieza:")
doc.add_paragraph(
    "• Eliminación de Duplicados - Complejidad: O(n)", style="List Bullet"
)
doc.add_paragraph("• Interpolación Lineal - Complejidad: O(n)", style="List Bullet")
doc.add_paragraph(
    "• Detección de Outliers mediante Z-Score - Complejidad: O(n)", style="List Bullet"
)

doc.add_heading("5.2 Algoritmos de Ordenamiento Implementados", 2)

table3 = doc.add_table(rows=13, cols=4)
table3.style = "Table Grid"
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = "#"
hdr_cells[1].text = "Algoritmo"
hdr_cells[2].text = "Complejidad"
hdr_cells[3].text = "Descripción"

algorithms = [
    ("1", "TimSort", "O(n log n)", "Combina Insertion Sort + Merge Sort"),
    ("2", "Comb Sort", "O(n²)", "Bubble Sort con gap decreciente"),
    ("3", "Selection Sort", "O(n²)", "Encuentra el mínimo en cada iteración"),
    ("4", "Tree Sort", "O(n log n)", "Usa BST para ordenar"),
    ("5", "Pigeonhole Sort", "O(n + k)", "Para rango(k) < n"),
    ("6", "Bucket Sort", "O(n + k)", "Distribuye en buckets"),
    ("7", "QuickSort", "O(n log n)", "Partición de Lomuto"),
    ("8", "HeapSort", "O(n log n)", "Garantiza O(n log n)"),
    ("9", "Bitonic Sort", "O(log² n)", "Para arquitecturas paralelas"),
    ("10", "Gnome Sort", "O(n²)", "Caminata por el array"),
    ("11", "Binary Insertion Sort", "O(n²)", "Búsqueda binaria para inserción"),
    ("12", "Radix Sort", "O(nk)", "Ordena por dígitos"),
]

for i, (num, alg, comp, desc) in enumerate(algorithms):
    row = table3.rows[i + 1].cells
    row[0].text = num
    row[1].text = alg
    row[2].text = comp
    row[3].text = desc

doc.add_heading("6. Resultados Obtenidos", 1)

doc.add_heading("6.1 Tabla de Tiempos de Ejecución", 2)
doc.add_paragraph(
    "Las mediciones fueron realizadas con un dataset de 12,792 registros:"
)

table4 = doc.add_table(rows=13, cols=3)
table4.style = "Table Grid"
hdr_cells = table4.rows[0].cells
hdr_cells[0].text = "Algoritmo"
hdr_cells[1].text = "Complejidad"
hdr_cells[2].text = "Tiempo (ms)"

times = [
    ("Radix Sort", "O(nk)", "3.14"),
    ("Pigeonhole Sort", "O(n + k)", "3.30"),
    ("Bucket Sort", "O(n + k)", "8.61"),
    ("TimSort", "O(n log n)", "33.47"),
    ("Comb Sort", "O(n²)", "95.10"),
    ("HeapSort", "O(n log n)", "107.71"),
    ("Bitonic Sort", "O(log² n)", "352.45"),
    ("Tree Sort", "O(n log n)", "1051.90"),
    ("QuickSort", "O(n log n)", "2107.53"),
    ("Binary Insertion Sort", "O(n²)", "4093.35"),
    ("Selection Sort", "O(n²)", "14090.68"),
    ("Gnome Sort", "O(n²)", "21215.64"),
]

for i, (alg, comp, t) in enumerate(times):
    row = table4.rows[i + 1].cells
    row[0].text = alg
    row[1].text = comp
    row[2].text = t

doc.add_heading("7. Conclusiones", 1)

doc.add_heading("7.1 Hallazgos Principales", 2)
doc.add_paragraph(
    "Del análisis de los resultados se pueden extraer varias conclusiones relevantes:"
)

doc.add_paragraph(
    "Complejidad teórica vs. Rendimiento real: Los algoritmos con mejor complejidad teórica no siempre son los más rápidos en la práctica. Radix Sort y Pigeonhole Sort superaron significativamente a algoritmos con O(n log n)."
)

doc.add_paragraph(
    "Eficiencia de TimSort: TimSort demostró ser muy eficiente (~33ms). Su capacidad de detectar runs naturales lo hace ideal para datos financieros."
)

doc.add_heading("8. Uso de Herramientas de Inteligencia Artificial", 1)
doc.add_paragraph("Se utilizó IA generativa como apoyo para:")
doc.add_paragraph(
    "• Estructuración del proyecto y generación de código base", style="List Bullet"
)
doc.add_paragraph("• Resolución de problemas técnicos", style="List Bullet")
doc.add_paragraph(
    "• El diseño algorítmico, análisis de complejidad y documentación fueron desarrollados de manera independiente.",
    style="List Bullet",
)

doc.add_heading("9. Ejecución del Proyecto", 1)
doc.add_paragraph("Comandos principales:")
doc.add_paragraph(
    "• task install - Crea entorno virtual e instala dependencias", style="List Bullet"
)
doc.add_paragraph("• task run - Ejecuta el pipeline completo", style="List Bullet")
doc.add_paragraph("• task run-full - Descarga datos reales", style="List Bullet")

doc.add_heading("10. Tecnologías Utilizadas", 1)
doc.add_paragraph("• Python 3.10+ - Lenguaje principal", style="List Bullet")
doc.add_paragraph("• requests - Peticiones HTTP", style="List Bullet")
doc.add_paragraph("• matplotlib - Visualizaciones", style="List Bullet")
doc.add_paragraph("• pandas - Manipulación de datos", style="List Bullet")
doc.add_paragraph("• Flask - API REST", style="List Bullet")
doc.add_paragraph("• Docker - Contenedores", style="List Bullet")

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.add_run(
    "Proyecto desarrollado para el curso de Análisis de Algoritmos - Universidad del Quindío - 2026-1"
).italic = True

doc.save(
    "E:\\Uniquindio\\Semestre 8\\Analisis de algoritmos\\investment-algorithm\\docs\\Informe_Seguimiento_1.docx"
)
print("Documento Word creado exitosamente!")
